#/n поиск в составе параграфа слова.. вроде как if 'Продукт' in row_text
#/n пересохранение или коллекция в таблицу
import docx
# pip install python-docx

# Открыть командную строку windows
# Установить pyinstaller

# pip install pyinstaller
#
# # Затем перейти в папку с Вашим файлом .py в командной строке (при помощи команды cd)
# # Запустить команду pyinstaller не забудьте указать имя вашего скрипта
#
# pyinstaller --onefile <your_script_name>.py
#
# # Всё - у вас в папке появится папка src и там будет .exe файл.

import os
import pandas  # +openpyxl
from tkinter import filedialog

paths = [] ## работа с файлами
#folder1 = os.getcwd() #?? свое
pathword = filedialog.askopenfilename()
folder = os.path.dirname(pathword)
print(folder)
txt=""
tt=0
f = open('text.txt', 'a', encoding='utf-8')
for root, dirs, files in os.walk(folder): # внутри него все папки и тп
    for file in files:
        #print(root, file)
        paths.append(os.path.join(root, file)) # собираем коллекцию
        txt = txt + os.path.join(root, file) + "\n"
        if tt == 1000:
            tt = 0
            f.write(txt)
            txt = ""
        # if file.endswith('docx') and not file.startswith('~'): # все докхс кроме .с тильды.
        #     paths.append(os.path.join(root, file)) # собираем коллекцию
        #     print(root, file)
# теперь получили коллекцию доков
# for path in paths:
#     doc = docx.Document(path)
#     properties = doc.core_properties
#     print('Автор документа:', properties.author)
#     print('Автор последней правки:', properties.last_modified_by)
#     print('Дата создания документа:', properties.created)
#     print('Дата последней правки:', properties.modified)
#     print('Дата последней печати:', properties.last_printed)
#     print('Количество сохранений:', properties.revision)
#     text = []  # заготавливаем лист

    # for table in doc.tables:  # перебираем ячейки в таблицах
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)

    # for paragraph in doc.paragraphs: # перебираем параграфы в тексте
    #     text.append(paragraph.text)  # сохраняем в лист чтобы потом посмотреть текст
    #     if paragraph.style.name == 'List Paragraph': #так например выглядит нумерованный список
    #         print(paragraph.text)
    #     for run in paragraph.runs:
    #         if run.italic:
    #             print(run.text)


            # print('Полужирный текст:', run.bold)
            # print('Подчёркнутый текст:', run.underline)
            # print('Зачёркнутый текст:', run.strike)
            # print('Название шрифта:', run.font.name)
            # print('Цвет текста, RGB:', run.font.color.rgb)
            # print('Цвет заливки текста:', run.font.highlight_color)
        # print('Выравнивание абзаца:', paragraph.alignment) #  None LEFT (0), center (1), RIGHT (2) или justify (3)
        # formatting = paragraph.paragraph_format
        # print('Отступ перед абзацем:', formatting.space_before)
        # print('Отступ после абзаца:', formatting.space_after)
        # print('Отступ слева:', formatting.left_indent)
        # print('Отступ справа:', formatting.right_indent)
        # print('Отступ первой строки абзаца:', formatting.first_line_indent.pt) в рзмере
        # print('Не отрывать от следующего абзаца:', formatting.keep_with_next)
        # print('Не разрывать абзац:', formatting.keep_together)
        # print('Абзац с новой страницы:', formatting.page_break_before)
        # print('Запрет висячих строк:', formatting.widow_control)

    #print('\n'.join(text))
f.write(txt)
print("wroten")
input()